from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.by import By
from constants import *
import time
from docx import Document
import requests

class Navigation():
    def __init__(self, teardown=False):
        self.teardown = teardown
        self.driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()))
        self.driver.maximize_window()
        self.driver.implicitly_wait(10)
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.teardown:
            self.quit()
    
    def get_url_navigator(self, url):
        try:
            self.driver.get(url)
            time.sleep(7)
        
        except Exception as e:
            print(f"Error {e}")
      
    def latest_news(self, xpath):
        news_bttn = self.driver.find_element(By.XPATH, xpath)
        news_bttn.click()
        time.sleep(5)
    
    def title_news(self, class_name, amount):
        try:
            elements = self.driver.find_elements(By.CSS_SELECTOR, class_name)
            titles = '\n\n'.join([f"{element.text}: {element.get_attribute('href')}" for element in elements[:amount]])
            print(titles)
            return titles
        except Exception as e:
            print(f"Error: {e}")
            
    def get_img(self, amount_img, class_img):
        links = self.driver.find_elements(By.XPATH, f'.//img[@class = "{class_img}"]')              
        
        for i, link in enumerate(links[:amount_img]):
            img_link = link.get_attribute('src')
            try:
                response = requests.get(img_link)
        
                if response == 200:
                    with open(f"{self.img_name}{i}.jpg", "wb") as file:
                        file.write(response.content)
                    print(f"Images downloaded successfully")
                else:
                    print(f"And error has ocurred. Status code: {response.status}")
            except Exception as e:
                print(f"Error in the download request {e}")
    
class WordDocumentCreator:
    def __init__(self, file_name):
        if type(file_name) != str:
            raise ValueError("The file name must be a string")
        self.file_name = file_name
        self.doc = Document()
        
    def add_paragraph(self, text):
        self.doc.add_paragraph(text)
    
    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level = level)
        
    def save_document(self):
        try:
            self.doc.save(self.file_name + ".docx")
            print(f"Word document {self.file_name} created and save whit success")
        except Exception as e:
            print(f"Error: {e}")
            
class image_management(Navigation):
    def __init__(self, driver, img_name):
        if type(img_name) != str:
            raise ValueError("The image name must be a string")
        super().__init__()
        self.img_name = img_name
        self.driver = driver

    def get_img(self, amount_img, class_img):
        links = self.driver.find_elements(By.XPATH, f'.//img[@class = "{class_img}"]')              
        
        for i, link in enumerate(links[:1+amount_img]):
            img_link = link.get_attribute('src')
            try:
                response = requests.get(img_link)
        
                if response.status_code == 200:
                    with open(f"{self.img_name}{i+1}.jpg", "wb") as file:
                        file.write(response.content)
                    print(f"Images downloaded successfully")
                else:
                    print(f"And error has ocurred. Status code: {response.status}")
            except Exception as e:
                print(f"Error in the download request {e}")       
                                
nav = Navigation()
#nav2 = image_management(nav.driver, "Image1")     
nav.get_url_navigator(dof["url"])
#nav.latest_news(universal["xpath"])
titles = nav.title_news(dof["title"], 1)
time.sleep(5)
word = WordDocumentCreator("Word1")
word.add_heading(f"News from dof")
word.add_paragraph(titles)
word.save_document()
#link1 = nav2.get_img(3, jornada["img"])
time.sleep(2)
        
        