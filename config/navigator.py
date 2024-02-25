from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.common.keys import Keys
import time
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document

#We created a class to use the webdriver options to navigate the page of the paper news
class Navigation(webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()))):
    def __init__(self,teardown=False):
        self.teardown = teardown
        super(Navigation, self).__init__()
        self.implicitly_wait(10)
    
    def __exit__(self, exc_type, exc_value, trace):
        if self.teardown:
            self.quit()
    
    def return_page(self):
        self.back()
        
    def get_url(self, page):
        try:
            self.get(page)
            self.implicitly_wait(10)
        except Exception as e:
            print(f"Can't access page {e}")
    
    def latest_news(self, xpath):
        news_bttn = self.find_element_by_xpath(xpath)
        print(news_bttn.text)
        news_bttn.click()
        
    def title_news(self, class_name, amount):
        try:
            elements = self.find_elements_by_css_selector(class_name)
            titles = '\n\n'.join([f"{element.text}: {element.get_attribute('href')}" for element in elements[:amount]])
            print(titles)
            return titles
        except Exception as e:
            print(f"Error: {e}")

class WordDocumentCreator:
    def __init__(self, file_name):
        self.file_name = file_name
        self.doc = Document()

    def add_paragraph(self, text):
        self.doc.add_paragraph(text)

    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)

    def save_document(self):
        self.doc.save(self.file_name)
        print(f"Documento Word '{self.file_name}' creado y guardado correctamente.")
